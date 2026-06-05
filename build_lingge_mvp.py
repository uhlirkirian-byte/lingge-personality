from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DOCX = "灵格40题MVP_v1.docx"
OUT_MD = "灵格40题MVP_v1.md"


SECTIONS = [
    {
        "title": "第一阶段：社会互动系统",
        "range": "Q1-Q10",
        "focus": "观察关系启动、社交进入、外界反馈敏感度、责任与控制倾向。",
    },
    {
        "title": "第二阶段：关系系统",
        "range": "Q11-Q20",
        "focus": "观察安全感、边界感、依恋模式、冲突修复、失望机制。",
    },
    {
        "title": "第三阶段：价值系统",
        "range": "Q21-Q30",
        "focus": "观察核心驱动力、自我价值来源、人生恐惧、长期追求。",
    },
    {
        "title": "第四阶段：压力与互搏系统",
        "range": "Q31-Q40",
        "focus": "观察压力反应、能量消耗、恢复方式、长期内耗与内部冲突。",
    },
]


QUESTIONS = [
    ("第一阶段：社会互动系统", "Q1", "哪类人突然发来一句“在吗？”，最容易让你立刻点开？", ["分开很久但曾经很重要的人", "久未联系但关系不错的朋友", "曾合作过的重要客户或领导", "家人或长期需要你负责的人"], "关系优先级 / 关系启动"),
    ("第一阶段：社会互动系统", "Q2", "点开消息前，你脑子里最先出现的念头更像：", ["他是不是遇到什么事了", "怎么突然想到我了", "这件事会不会影响我接下来安排", "先看内容，不急着判断"], "关注点排序 / 安全感"),
    ("第一阶段：社会互动系统", "Q3", "一个你很在意的人误会你，而且暂时不愿听解释，你更可能：", ["尽快把事实说清楚", "表面沉默，心里反复委屈", "情绪上来，忍不住反击", "不再解释，慢慢拉开距离"], "误会处理 / 表达与压抑"),
    ("第一阶段：社会互动系统", "Q4", "如果明天突然不用工作，也不用对任何人负责，你今晚更像：", ["立刻计划明天怎么安排", "照常生活，顺其自然", "做点平时没空做的事", "一时不知道自己真正想做什么"], "控制感 / 自由感"),
    ("第一阶段：社会互动系统", "Q5", "被朋友邀请参加一个陌生聚会，刚到现场时你通常会：", ["主动开启聊天", "先待在一旁，等别人来搭话", "观察一会儿，找舒服的人交流", "寻找自然时机早点离开"], "社交进入方式"),
    ("第一阶段：社会互动系统", "Q6", "一个能力和你差不多的同事突然成了你的上司，你第一反应更像：", ["接受现实，先看他怎么做", "心里不服，但表面配合", "反思自己差在哪里", "开始评估自己是否还适合这里"], "竞争反应 / 自我价值"),
    ("第一阶段：社会互动系统", "Q7", "你最难长期相处的人，更像哪一种？", ["说话不算数的人", "情绪失控的人", "控制欲强的人", "只顾自己的人"], "关系底线 / 触发点"),
    ("第一阶段：社会互动系统", "Q8", "你主动给一个很在意的人发消息，对方迟迟没回，你更像：", ["觉得他只是忙", "反复看手机", "开始想是不是自己说错话", "告诉自己别太在意，但心里仍受影响"], "关系安全感锚点"),
    ("第一阶段：社会互动系统", "Q9", "小时候，大人最常评价你的方式更接近：", ["懂事，让人省心", "有主见，不太听话", "聪明，但难管", "安静，不太表达自己"], "早期角色 / 责任感"),
    ("第一阶段：社会互动系统", "Q10", "如果人生能重来一次，你最想改变的是：", ["错过的人", "做错的决定", "没坚持下去的梦想", "我不会改太多，只想看清当时的自己"], "人生遗憾入口"),
    ("第二阶段：关系系统", "Q11", "一个你很在意的人明显没有以前热情了，你更可能：", ["主动问是不是发生了什么", "表面正常，心里受影响但不表露", "观察一段时间，再决定要不要继续投入", "也开始减少投入"], "关系降温反应"),
    ("第二阶段：关系系统", "Q12", "和亲近的人发生矛盾后，你更常见的状态是：", ["希望尽快说开，不喜欢拖着", "需要时间冷静，之后仍愿意沟通", "嘴上说没事，其实一直记着", "用距离表达不满，甚至不想沟通"], "冲突修复方式"),
    ("第二阶段：关系系统", "Q13", "一个平时几乎不联系的人，总在需要帮忙时才出现，你通常会：", ["只要帮得上，还是会帮", "心里明白，但表面维持客气", "两三次之后直接减少回应", "表现得敷衍，让对方知难而退"], "边界感 / 讨好倾向"),
    ("第二阶段：关系系统", "Q14", "一群人聊天时，你说的话被大家略过了，你更像：", ["无所谓，继续参与", "安静下来，减少表达", "表面自然，心里有点失落", "找机会重新把话题接回来"], "表达需求 / 被看见"),
    ("第二阶段：关系系统", "Q15", "当你开始喜欢一个人时，你通常更像：", ["主动靠近，释放好感", "先观察对方有没有兴趣", "希望对方先主动", "制造接触机会，但不明说"], "靠近策略"),
    ("第二阶段：关系系统", "Q16", "别人真诚夸奖你时，你更容易：", ["自然接受，并开心一下", "谦虚带过，不太当回事", "怀疑对方只是客气", "表面接受，心里有点别扭"], "自我价值接收方式"),
    ("第二阶段：关系系统", "Q17", "为了维持一段重要关系，你通常能接受委屈自己到什么程度？", ["基本不委屈自己", "小事可以，大事不行", "常常先让着对方", "当下忍着，之后会爆发"], "委屈阈值 / 责任感"),
    ("第二阶段：关系系统", "Q18", "当你在意的人和别人明显走得很近时，你更像：", ["直接表达自己的在意", "表面不动声色，但持续观察", "心里不舒服，但选择忍着不说", "尊重对方选择，不主动干涉"], "关系竞争机制锚点"),
    ("第二阶段：关系系统", "Q19", "一个人连续几次让你失望后，你通常会：", ["找机会认真谈一次", "继续给机会，但心里划出最后红线", "表面正常，心里已经把对方移出重要位置", "直接抽离，不再投入"], "失望机制"),
    ("第二阶段：关系系统", "Q20", "在亲密关系里，你最怕的是：", ["被忽视", "被控制", "被背叛", "不被理解"], "核心关系恐惧锚点"),
    ("第三阶段：价值系统", "Q21", "如果未来十年只能稳定得到一种结果，你更想要：", ["稳定富足的生活", "做成一件自己真正认同的事", "被很多人认可与欣赏", "和重要的人长期在一起"], "人生终极追求"),
    ("第三阶段：价值系统", "Q22", "当你想给重要的人留下好印象时，你最自然展现的是：", ["舒服、好相处", "可靠、有分寸", "有趣、有魅力", "有想法、有特点"], "社会展示面"),
    ("第三阶段：价值系统", "Q23", "如果突然拥有一大笔钱，你最先想解决的是：", ["现实压力和安全感", "时间自由，不再被困住", "证明自己有能力做到", "体验一直想过的人生"], "财富意义 / 控制感"),
    ("第三阶段：价值系统", "Q24", "你更怕哪种人生状态？", ["一直很忙，却不知道为了什么", "很稳定，但毫无波澜", "有梦想，却始终做不成", "被很多人需要，却没人真正懂你"], "人生核心恐惧锚点"),
    ("第三阶段：价值系统", "Q25", "你最容易羡慕哪类人？", ["活得轻松自在的人", "很有能力的人", "被很多人喜欢的人", "清楚知道自己要什么的人"], "投射羡慕对象"),
    ("第三阶段：价值系统", "Q26", "一个你并不认可的人轻描淡写说你“也就那样”，你更像：", ["当下不爽，想反击", "表面无事，但会记很久", "觉得无所谓，他说他的", "开始怀疑自己是不是真的不够好"], "否定触发 / 自我价值"),
    ("第三阶段：价值系统", "Q27", "你努力做事时，最深层动力更像：", ["不想再过被动的人生", "想让重要的人看得起我", "想证明自己的价值", "想活得舒服、自洽一点"], "核心驱动力锚点"),
    ("第三阶段：价值系统", "Q28", "认真投入一件事三个月，却没人看见也没回报，你通常会：", ["继续做，我认同它就够了", "会动摇，但还能坚持一阵", "开始怀疑值不值得", "转向更有反馈的方向"], "回报需求 / 意义感"),
    ("第三阶段：价值系统", "Q29", "你最难接受自己变成哪种人？", ["对生活失去感受力的人", "遇事总逃避责任的人", "明明有能力却一直没做成事的人", "一直按别人标准活着的人"], "自我否定底线"),
    ("第三阶段：价值系统", "Q30", "别人私下提起你时，你最希望听到：", ["这人靠谱，能托付事", "这人真有本事", "这人挺有意思，不一样", "这人活得明白"], "理想评价"),
    ("第四阶段：压力与互搏系统", "Q31", "当压力连续累积一段时间，你最明显的变化通常是：", ["压力越大，反而越有战斗力", "容易烦躁，对人没耐心", "表面正常，心里越来越累", "开始拖着，不想面对"], "压力模式锚点"),
    ("第四阶段：压力与互搏系统", "Q32", "遇到一时解决不了的压力，到了晚上你更像：", ["先放一放，照样睡觉", "脑子停不下来，很难睡着", "靠酒精、药物或其他方式帮助入睡", "想找人聊一聊，缓过来再睡"], "压力释放"),
    ("第四阶段：压力与互搏系统", "Q33", "当你心里难受时，第一反应更常是：", ["找人说出来", "自己慢慢消化", "做点别的事分散注意力", "装作没事，继续过日子"], "情绪处理 / 防御机制"),
    ("第四阶段：压力与互搏系统", "Q34", "你最容易在哪种情况下情绪一下子上来？", ["被人当面否定", "重要成果被人抢功", "被强迫做不愿做的事", "被信任的人当场伤害"], "情绪触发点"),
    ("第四阶段：压力与互搏系统", "Q35", "当生活一段时间都很顺时，你更容易：", ["状态更好，做事更积极", "放松下来，多享受生活", "心里反而有点不安", "开始追求更大的目标"], "顺境反应"),
    ("第四阶段：压力与互搏系统", "Q36", "当你状态低到提不起劲时，通常最先靠什么恢复？", ["一个人待着，安静充电", "找懂你的人聊聊", "做成一件小事找回节奏", "换个新环境或新刺激重新启动"], "恢复机制"),
    ("第四阶段：压力与互搏系统", "Q37", "你最常把能量耗在哪件事上？", ["想太多，还没开始先累了", "对别人太上心", "拖着没行动", "明知不值得，还继续投入"], "能量黑洞锚点"),
    ("第四阶段：压力与互搏系统", "Q38", "一件重要的事失败后，你通常更像：", ["很快调整，继续往前", "低落几天，再慢慢恢复", "反复想很久，走不出来", "看起来恢复了，心里一直没过去"], "失败恢复"),
    ("第四阶段：压力与互搏系统", "Q39", "如果长期处在总被否定的环境里，你更容易变成：", ["更想证明自己", "表面硬撑，心里越来越压抑", "慢慢怀疑自己，失去动力", "情感抽离，不再在乎评价"], "长期否定反应"),
    ("第四阶段：压力与互搏系统", "Q40", "当你面对生活或工作中的无力感时，最难接受的是：", ["对很多事慢慢失去热情", "努力了也看不到希望", "渐渐觉得自己不重要", "身边始终没人真正理解你"], "无力感来源"),
]


ANCHORS = [
    ("Q8", "关系安全感", "等待焦虑、确认需求、关系敏感度", "Q11 / Q18 / Q20"),
    ("Q18", "关系竞争机制", "占有感、边界感、情感竞争反应", "Q8 / Q11 / Q20"),
    ("Q20", "核心关系恐惧", "被忽视、被控制、被背叛、不被理解", "Q8 / Q18 / Q17"),
    ("Q24", "人生核心恐惧", "意义、稳定、成就、被理解", "Q21 / Q23 / Q27 / Q29"),
    ("Q27", "核心驱动力", "为什么努力、为什么投入", "Q21 / Q24 / Q28 / Q30"),
    ("Q31", "压力模式", "高压状态下的真实反应", "Q32 / Q33 / Q37 / Q38"),
    ("Q37", "能量损耗来源", "长期内耗模式", "Q31 / Q35 / Q36 / Q40"),
]


GROUPS = [
    ("安全感组", "Q8 / Q11 / Q18 / Q20"),
    ("自我价值组", "Q16 / Q22 / Q26 / Q27 / Q30"),
    ("责任感组", "Q9 / Q17 / Q22 / Q29 / Q30"),
    ("控制感组", "Q4 / Q6 / Q23 / Q31"),
    ("意义感组", "Q21 / Q24 / Q27 / Q28 / Q29"),
    ("内耗组", "Q31 / Q37 / Q38 / Q40"),
    ("靠近 VS 防御", "Q8 / Q11 / Q15 / Q18 / Q20"),
    ("成长 VS 安全", "Q21 / Q23 / Q24 / Q27"),
    ("表达 VS 压抑", "Q3 / Q12 / Q14 / Q33"),
    ("自由 VS 责任", "Q4 / Q21 / Q24 / Q29 / Q30"),
    ("认可 VS 自我实现", "Q22 / Q25 / Q27 / Q30"),
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width
            row.cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(row.cells[idx])


def add_option_paragraph(cell, text):
    p = cell.paragraphs[0] if not cell.paragraphs[0].text else cell.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.1
    run = p.add_run(text)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(9.5)


def add_question_table(doc, qid, question, options, target):
    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    set_table_widths(table, [Inches(0.65), Inches(5.85)])
    table.cell(0, 0).merge(table.cell(1, 0))
    set_cell_shading(table.cell(0, 0), "F2F4F7")
    qcell = table.cell(0, 0)
    p = qcell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(qid)
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Arial"
    table.cell(0, 1).text = ""
    p = table.cell(0, 1).paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(question)
    run.bold = True
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(10.5)
    opt_cell = table.cell(1, 1)
    opt_cell.text = ""
    labels = ["A", "B", "C", "D"]
    for label, opt in zip(labels, options):
        add_option_paragraph(opt_cell, f"□ {label}. {opt}")
    meta = opt_cell.add_paragraph()
    meta.paragraph_format.space_before = Pt(2)
    meta.paragraph_format.space_after = Pt(0)
    mrun = meta.add_run(f"观察窗口：{target}")
    mrun.font.name = "Microsoft YaHei"
    mrun._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    mrun.font.size = Pt(8.5)
    mrun.font.color.rgb = RGBColor(85, 85, 85)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def style_doc(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)
    for name, size, color in [
        ("Heading 1", 18, "000000"),
        ("Heading 2", 14, "000000"),
        ("Heading 3", 12, "434343"),
    ]:
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name != "Heading 1"


def add_simple_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_widths(table, widths)
    for i, header in enumerate(headers):
        set_cell_shading(table.cell(0, i), "F2F4F7")
        p = table.cell(0, i).paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(9.5)
    for row_data in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_data):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(value)
            run.font.name = "Microsoft YaHei"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            run.font.size = Pt(9)
            set_cell_margins(cells[i])
    doc.add_paragraph()


def build_docx():
    doc = Document()
    style_doc(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    tr = title.add_run("灵格 40 题人格结构采样 MVP")
    tr.font.name = "Microsoft YaHei"
    tr._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    tr.font.size = Pt(24)
    tr.bold = True

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    sr = subtitle.add_run("V1.0 | 行为投射题 / 六维人格结构 / 互搏识别优先")
    sr.font.name = "Microsoft YaHei"
    sr._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    sr.font.size = Pt(10.5)
    sr.font.color.rgb = RGBColor(85, 85, 85)

    doc.add_heading("MVP 定位", level=1)
    for text in [
        "这不是人格分类题，也不是心理咨询诊断。它是一套 3-5 分钟的人格结构采样器，用 40 个行为场景观察用户在关系、价值、压力与内部冲突中的稳定选择模式。",
        "分析时禁止单题下结论。至少三个关联题共同指向同一结构，才进入报告判断。报告目标不是告诉用户“你是什么人”，而是解释“你为什么总这样”。",
    ]:
        doc.add_paragraph(text)

    doc.add_heading("四阶段结构", level=1)
    add_simple_table(
        doc,
        ["阶段", "题号", "观察重点"],
        [(s["title"], s["range"], s["focus"]) for s in SECTIONS],
        [Inches(1.8), Inches(1.0), Inches(4.7)],
    )

    doc.add_heading("答题说明", level=1)
    for text in [
        "每题选择最接近你真实第一反应的选项，不需要选择看起来更正确或更成熟的答案。",
        "若两个选项都像自己，优先选择压力更大、关系更重要、代价更高时更常出现的那个反应。",
        "本 MVP 建议先收集完整答案、用户年龄段、关系状态、当前最大困扰，以及用户读完报告后的反馈。反馈优先级高于题目本身。",
    ]:
        doc.add_paragraph(text, style=None)

    doc.add_heading("40 题正式问卷", level=1)
    current = None
    for section, qid, question, options, target in QUESTIONS:
        if section != current:
            current = section
            doc.add_heading(section, level=2)
        add_question_table(doc, qid, question, options, target)

    doc.add_section(WD_SECTION_START.NEW_PAGE)
    doc.add_heading("分析使用说明", level=1)
    doc.add_heading("超级锚点题", level=2)
    add_simple_table(
        doc,
        ["题号", "锚点", "测量内容", "关联题"],
        ANCHORS,
        [Inches(0.7), Inches(1.35), Inches(3.1), Inches(2.35)],
    )
    doc.add_heading("交叉验证组", level=2)
    add_simple_table(
        doc,
        ["结构组", "关联题"],
        GROUPS,
        [Inches(2.2), Inches(5.3)],
    )
    doc.add_heading("报告生成原则", level=2)
    for text in [
        "优先级：发现冲突 > 发现盲区 > 发现隐藏动机 > 发现长期循环 > 描述人格特点。",
        "高价值报告必须指出代价。例如：不是“你很重感情”，而是“你最难放下的往往不是人，而是自己已经投入过的时间、期待和解释权”。",
        "典型报告结构：人物标题、核心行为模式、真正驱动力、主要防御机制、内部互搏、长期循环、一个最值得观察的成长方向。",
    ]:
        doc.add_paragraph(text)
    doc.add_heading("MVP 成功标准", level=2)
    for text in [
        "用户愿意完成测试。",
        "用户愿意读完整份报告。",
        "用户反馈“这个我以前没这样想过”。",
        "用户愿意保存、分享或拿报告和朋友讨论。",
        "样本反馈能反向帮助修正题目与映射规则。",
    ]:
        p = doc.add_paragraph(style=None)
        p.style = doc.styles["Normal"]
        p.add_run("□ ").bold = True
        p.add_run(text)

    doc.save(OUT_DOCX)


def build_markdown():
    lines = [
        "# 灵格 40 题人格结构采样 MVP",
        "",
        "V1.0 | 行为投射题 / 六维人格结构 / 互搏识别优先",
        "",
        "## MVP 定位",
        "",
        "这不是人格分类题，也不是心理咨询诊断。它是一套 3-5 分钟的人格结构采样器，用 40 个行为场景观察用户在关系、价值、压力与内部冲突中的稳定选择模式。",
        "",
        "分析时禁止单题下结论。至少三个关联题共同指向同一结构，才进入报告判断。报告目标不是告诉用户“你是什么人”，而是解释“你为什么总这样”。",
        "",
        "## 四阶段结构",
        "",
    ]
    for s in SECTIONS:
        lines.append(f"- {s['title']}（{s['range']}）：{s['focus']}")
    lines.extend(["", "## 答题说明", ""])
    for item in [
        "每题选择最接近你真实第一反应的选项，不需要选择看起来更正确或更成熟的答案。",
        "若两个选项都像自己，优先选择压力更大、关系更重要、代价更高时更常出现的那个反应。",
        "本 MVP 建议先收集完整答案、用户年龄段、关系状态、当前最大困扰，以及用户读完报告后的反馈。",
    ]:
        lines.append(f"- {item}")
    lines.extend(["", "## 40 题正式问卷", ""])
    current = None
    for section, qid, question, options, target in QUESTIONS:
        if section != current:
            current = section
            lines.extend(["", f"### {section}", ""])
        lines.append(f"**{qid}. {question}**")
        for label, option in zip(["A", "B", "C", "D"], options):
            lines.append(f"- {label}. {option}")
        lines.append(f"- 观察窗口：{target}")
        lines.append("")
    lines.extend(["## 分析使用说明", "", "### 超级锚点题", ""])
    for row in ANCHORS:
        lines.append(f"- {row[0]} | {row[1]} | {row[2]} | 关联：{row[3]}")
    lines.extend(["", "### 交叉验证组", ""])
    for row in GROUPS:
        lines.append(f"- {row[0]}：{row[1]}")
    lines.extend(["", "### 报告生成原则", ""])
    for item in [
        "优先级：发现冲突 > 发现盲区 > 发现隐藏动机 > 发现长期循环 > 描述人格特点。",
        "高价值报告必须指出代价，而不是只描述特点。",
        "典型报告结构：人物标题、核心行为模式、真正驱动力、主要防御机制、内部互搏、长期循环、一个最值得观察的成长方向。",
    ]:
        lines.append(f"- {item}")
    with open(OUT_MD, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))


if __name__ == "__main__":
    build_docx()
    build_markdown()
